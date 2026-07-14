"""
Valutazione qualitativa del contributo di Neo4j (GraphRAG).

Per un campione di 20 domande, genera le risposte complete con Mistral
in due modalità:
  - senza arricchimento grafo (usa_grafo=False)
  - con arricchimento grafo  (usa_grafo=True)

Salva le risposte affiancate in un file JSON e in un documento Word
per la valutazione manuale da parte del ricercatore.

Metriche automatiche calcolate:
  - lunghezza della risposta (numero di caratteri)
  - presenza di citazioni di leggi modificanti
  - presenza di riferimenti a versioni storiche
  - BERTScore (se disponibile) per similarità tra le due risposte

USO:
    export NEO4J_PASS="tesi2026"
    export OS_PASS="PasswordForte123"
    export COHERE_API_KEY="..."
    python3 valutazione_neo4j.py
"""

import os, sys, json, time, importlib.util
from pathlib import Path
from datetime import datetime

# Configurazione
BASE_DIR   = Path(__file__).parent
CHATBOT_PY = BASE_DIR / "chatbot.py"
OUTPUT_DIR = BASE_DIR / "risultati_neo4j"
OUTPUT_DIR.mkdir(exist_ok=True)

# Campione di 20 domande — mix di categorie e tipologie
# Scelte per massimizzare il valore informativo del confronto con/senza grafo:
# domande su norme modificate (dove il grafo aggiunge storia delle modifiche),
# domande con articoli che citano altre norme (dove il grafo aggiunge relazioni)
DOMANDE_CAMPIONE = [
    # Categoria 1 — con numero e codice (il grafo aggiunge storia modifiche)
    "Cosa prevede l'articolo 2043 del Codice Civile?",
    "Cosa stabilisce l'articolo 575 del Codice Penale?",
    "Come è disciplinato l'articolo 116 del Nuovo Codice della Strada?",
    "Cosa prevede l'articolo 380 del Codice di Procedura Penale?",
    "Cosa disciplina l'articolo 275 del Codice di Procedura Penale?",

    # Categoria 2 — semantica con codice (il grafo aggiunge contesto relazionale)
    "Cosa prevede il Codice della Strada sull'uso del telefono alla guida?",
    "Quali sono le sanzioni per abbandono di rifiuti secondo il Codice dell'Ambiente?",
    "Come disciplina il Codice delle Assicurazioni il risarcimento diretto?",
    "Cosa prevede il Codice della Privacy sulle comunicazioni commerciali?",
    "Come regolamenta il Codice del Consumo il diritto di recesso?",

    # Categoria 3 — temporale (il grafo è più utile per ricostruire la storia)
    "Cosa prevedeva l'articolo 116 del Codice della Strada nel 1993?",
    "Come era disciplinato l'articolo 275 del Codice di Procedura Penale nel 1992?",
    "Qual era la disciplina dell'articolo 1284 del Codice Civile nel 1995?",

    # Categoria 4 — senza codice né numero (il grafo aiuta il contesto)
    "Ho preso una multa per eccesso di velocità, quanti punti mi tolgono?",
    "Posso restituire un prodotto comprato online senza dare spiegazioni?",
    "Cosa rischio se abbandono i rifiuti in campagna?",

    # Categoria 5 — evolutive (il grafo è più utile qui)
    "Come è cambiato nel tempo l'articolo 380 del Codice di Procedura Penale?",
    "Come si è evoluta la disciplina dell'articolo 116 del Codice della Strada?",
    "Quante versioni ha avuto l'articolo 1284 del Codice Civile sugli interessi?",
    "Come è cambiata nel tempo la disciplina dell'articolo 275 del c.p.p. sulle misure cautelari?",
]

# Caricamento chatbot
def carica_chatbot():
    os.environ["TOP_K"]        = "10"
    os.environ["USE_HYDE"]     = "0"
    os.environ["USE_KW_EXP"]   = "0"
    os.environ["SEARCH_SIZE"]  = "8"
    spec   = importlib.util.spec_from_file_location("chatbot", CHATBOT_PY)
    modulo = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(modulo)
    return modulo

# Metriche automatiche
def analizza_risposta(risposta: str) -> dict:
    """Estrae metriche automatiche dalla risposta generata."""
    r = risposta.lower()
    return {
        "lunghezza_caratteri": len(risposta),
        "n_parole": len(risposta.split()),
        "cita_leggi_modificanti": any(w in r for w in
            ["modificat", "sostitu", "abrocat", "novell", "riformat"]),
        "cita_versioni_storiche": any(w in r for w in
            ["vigente dal", "fino al", "versione", "storico", "precedente"]),
        "cita_articoli_collegati": risposta.count("art.") + risposta.count("Art."),
        "contiene_fallback": "FONTI_NON_SUFFICIENTI" in risposta,
        "contiene_avviso_allucinazione": "⚠" in risposta or "verificare" in r,
    }

def bertscore_similarita(testo1: str, testo2: str) -> float:
    """Calcola BERTScore tra due testi se bert_score è disponibile."""
    try:
        from bert_score import score
        P, R, F = score([testo1], [testo2], lang="it", verbose=False)
        return round(float(F[0]), 4)
    except ImportError:
        return None

# Valutazione singola domanda
def valuta_domanda(chatbot, domanda: str, idx: int, totale: int) -> dict:
    print(f"\n[{idx}/{totale}] {domanda[:65]}...")

    risultato = {
        "domanda": domanda,
        "timestamp": datetime.now().isoformat(),
        "senza_grafo": {},
        "con_grafo": {},
        "confronto": {},
    }

    # ── Senza grafo ────────────────────────────────────────────────────────────
    print("  Generazione senza grafo...", end=" ", flush=True)
    t0 = time.time()
    try:
        risposta_no = chatbot.chatbot(domanda, usa_grafo=False)
        t_no = round(time.time() - t0, 1)
        risultato["senza_grafo"] = {
            "risposta": risposta_no,
            "tempo_secondi": t_no,
            "metriche": analizza_risposta(risposta_no),
        }
        print(f"OK ({t_no}s)")
    except Exception as e:
        print(f"ERRORE: {e}")
        risultato["senza_grafo"] = {"errore": str(e)}

    # ── Con grafo ──────────────────────────────────────────────────────────────
    print("  Generazione con grafo...", end=" ", flush=True)
    t0 = time.time()
    try:
        risposta_si = chatbot.chatbot(domanda, usa_grafo=True)
        t_si = round(time.time() - t0, 1)
        risultato["con_grafo"] = {
            "risposta": risposta_si,
            "tempo_secondi": t_si,
            "metriche": analizza_risposta(risposta_si),
        }
        print(f"OK ({t_si}s)")
    except Exception as e:
        print(f"ERRORE: {e}")
        risultato["con_grafo"] = {"errore": str(e)}

    # ── Confronto automatico ───────────────────────────────────────────────────
    if "risposta" in risultato["senza_grafo"] and "risposta" in risultato["con_grafo"]:
        r_no = risultato["senza_grafo"]["risposta"]
        r_si = risultato["con_grafo"]["risposta"]
        bs   = bertscore_similarita(r_no, r_si)
        m_no = risultato["senza_grafo"]["metriche"]
        m_si = risultato["con_grafo"]["metriche"]
        risultato["confronto"] = {
            "bertscore_f1": bs,
            "delta_lunghezza": m_si["lunghezza_caratteri"] - m_no["lunghezza_caratteri"],
            "delta_parole": m_si["n_parole"] - m_no["n_parole"],
            "grafo_aggiunge_leggi_modificanti":
                m_si["cita_leggi_modificanti"] and not m_no["cita_leggi_modificanti"],
            "grafo_aggiunge_versioni_storiche":
                m_si["cita_versioni_storiche"] and not m_no["cita_versioni_storiche"],
            "grafo_aggiunge_articoli_collegati":
                m_si["cita_articoli_collegati"] > m_no["cita_articoli_collegati"],
            "delta_tempo": round(risultato["con_grafo"]["tempo_secondi"]
                                 - risultato["senza_grafo"]["tempo_secondi"], 1),
        }

    return risultato

# Salvataggio Word per valutazione manuale
def salva_word(risultati: list, output_path: Path):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("python-docx non installato — salto il file Word.")
        return

    doc = Document()
    doc.add_heading("Valutazione qualitativa: risposta con e senza Neo4j", 0)
    doc.add_paragraph(
        "Per ogni domanda valuta su scala 1-5: "
        "Completezza | Accuratezza | Contesto storico | Citazione fonti | Utilità complessiva"
    )

    for i, r in enumerate(risultati, 1):
        doc.add_heading(f"{i}. {r['domanda']}", level=1)

        # Senza grafo
        doc.add_heading("Risposta SENZA Neo4j", level=2)
        p = doc.add_paragraph(r.get("senza_grafo", {}).get("risposta", "ERRORE"))
        m = r.get("senza_grafo", {}).get("metriche", {})
        doc.add_paragraph(
            f"Tempo: {r.get('senza_grafo',{}).get('tempo_secondi','?')}s | "
            f"Parole: {m.get('n_parole','?')} | "
            f"Cita modifiche: {'Sì' if m.get('cita_leggi_modificanti') else 'No'}"
        ).italic = True

        # Griglia valutazione manuale
        doc.add_paragraph("Valutazione (1-5): Completezza ___ | Accuratezza ___ | "
                          "Contesto storico ___ | Citazione fonti ___ | Utilità ___")

        # Con grafo
        doc.add_heading("Risposta CON Neo4j", level=2)
        doc.add_paragraph(r.get("con_grafo", {}).get("risposta", "ERRORE"))
        m2 = r.get("con_grafo", {}).get("metriche", {})
        doc.add_paragraph(
            f"Tempo: {r.get('con_grafo',{}).get('tempo_secondi','?')}s | "
            f"Parole: {m2.get('n_parole','?')} | "
            f"Cita modifiche: {'Sì' if m2.get('cita_leggi_modificanti') else 'No'}"
        ).italic = True

        doc.add_paragraph("Valutazione (1-5): Completezza ___ | Accuratezza ___ | "
                          "Contesto storico ___ | Citazione fonti ___ | Utilità ___")

        # Confronto automatico
        c = r.get("confronto", {})
        if c:
            doc.add_paragraph(
                f"Confronto automatico — "
                f"Δ parole: {c.get('delta_parole','?'):+d} | "
                f"Δ tempo: {c.get('delta_tempo','?'):+.1f}s | "
                f"BERTScore: {c.get('bertscore_f1','N/D')} | "
                f"Grafo aggiunge storia: {'Sì' if c.get('grafo_aggiunge_versioni_storiche') else 'No'} | "
                f"Grafo aggiunge leggi mod.: {'Sì' if c.get('grafo_aggiunge_leggi_modificanti') else 'No'}"
            ).bold = True

        doc.add_paragraph("─" * 80)

    doc.save(str(output_path))
    print(f"File Word salvato: {output_path}")

# Main
def main():
    if not CHATBOT_PY.exists():
        print(f"ERRORE: {CHATBOT_PY} non trovato.")
        sys.exit(1)

    if not os.getenv("NEO4J_PASS"):
        print("ATTENZIONE: NEO4J_PASS non impostata.")
        print("  export NEO4J_PASS='tesi2026'")

    print("Caricamento chatbot.py...")
    chatbot = carica_chatbot()

    print(f"\nAvvio valutazione su {len(DOMANDE_CAMPIONE)} domande...")
    risultati = []

    for i, domanda in enumerate(DOMANDE_CAMPIONE, 1):
        r = valuta_domanda(chatbot, domanda, i, len(DOMANDE_CAMPIONE))
        risultati.append(r)

        # Salva progressivamente
        json_path = OUTPUT_DIR / "risultati_neo4j.json"
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(risultati, f, indent=2, ensure_ascii=False)

    # Riepilogo automatico
    print("\n" + "="*60)
    print("RIEPILOGO CONFRONTO AUTOMATICO")
    print("="*60)
    confronti = [r["confronto"] for r in risultati if r.get("confronto")]
    if confronti:
        n = len(confronti)
        print(f"Domande confrontate: {n}")
        print(f"Δ parole medio (grafo - no grafo): "
              f"{sum(c.get('delta_parole',0) for c in confronti)/n:+.1f}")
        print(f"Δ tempo medio (grafo - no grafo): "
              f"{sum(c.get('delta_tempo',0) for c in confronti)/n:+.1f}s")
        print(f"Grafo aggiunge leggi modificanti: "
              f"{sum(1 for c in confronti if c.get('grafo_aggiunge_leggi_modificanti'))}/{n}")
        print(f"Grafo aggiunge versioni storiche: "
              f"{sum(1 for c in confronti if c.get('grafo_aggiunge_versioni_storiche'))}/{n}")
        bs_values = [c["bertscore_f1"] for c in confronti if c.get("bertscore_f1")]
        if bs_values:
            print(f"BERTScore medio: {sum(bs_values)/len(bs_values):.4f}")

    # Salva Word per valutazione manuale
    word_path = OUTPUT_DIR / "valutazione_manuale_neo4j.docx"
    salva_word(risultati, word_path)

    print(f"\nJSON: {json_path}")
    print(f"Word: {word_path}")

if __name__ == "__main__":
    main()
